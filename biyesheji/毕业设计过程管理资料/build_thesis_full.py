# -*- coding: utf-8 -*-
"""Generate full thesis docx + figure PNGs. Run: python build_thesis_full.py"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

from thesis_zh_body import (
    关键词行,
    附录段落,
    结论段落,
    第1章,
    第2章,
    第3章,
    第4章,
    第5章,
    摘要段落,
    致谢段落,
)

BASE = Path(__file__).resolve().parent
FIG_DIR = BASE / "论文插图"
OUT_DOCX = BASE / "长沙学院毕业设计_B20220302130_谢志坚_全稿_v2.docx"


def set_doc_default_font(doc, font_name="宋体"):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_body(doc, text, first_indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(0)


def add_fig_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def code_lines_to_png(lines, out_path, title="", font_size=13):
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    pad = 16
    line_h = font_size + 6
    w = 920
    h = pad * 2 + (len(lines) + (2 if title else 0)) * line_h + 20
    img = Image.new("RGB", (w, h), (255, 255, 255))
    dr = ImageDraw.Draw(img)
    try:
        fmono = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", font_size)
    except OSError:
        fmono = ImageFont.load_default()
    try:
        fsim = ImageFont.truetype("C:/Windows/Fonts/simfang.ttf", font_size)
    except OSError:
        fsim = fmono
    y = pad
    if title:
        dr.text((pad, y), title, fill=(0, 0, 120), font=fsim)
        y += line_h + 4
    for line in lines:
        dr.text((pad, y), line, fill=(0, 0, 0), font=fmono)
        y += line_h
    img.save(out_path, "PNG")


def draw_block_diagram(out_path):
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    w, h = 820, 420
    img = Image.new("RGB", (w, h), (255, 255, 255))
    dr = ImageDraw.Draw(img)
    try:
        f = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 16)
        fs = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 14)
    except OSError:
        f = fs = ImageFont.load_default()
    boxes = [
        (80, 40, 200, 90, "STM32F411"),
        (320, 40, 480, 90, "显示+触摸"),
        (520, 40, 700, 90, "传感器组"),
        (200, 160, 380, 210, "FreeRTOS"),
        (420, 160, 600, 210, "LVGL 界面"),
        (260, 280, 540, 330, "电源/充电/蓝牙"),
    ]
    for x1, y1, x2, y2, t in boxes:
        dr.rectangle([x1, y1, x2, y2], outline=(0, 0, 0), width=2)
        l, top, r, bot = dr.textbbox((0, 0), t, font=fs)
        tw, th = r - l, bot - top
        dr.text((x1 + (x2 - x1 - tw) // 2, y1 + (y2 - y1 - th) // 2), t, fill=0, font=fs)
    dr.text((260, 8), "图2-1  系统总体框图（示意）", fill=(80, 80, 80), font=f)
    img.save(out_path, "PNG")


def build_figures():
    hr_lines = [
        "/* heart rate: peak spacing */",
        "uint16_t HR_Calculate(uint16_t present_dat,uint32_t present_time)",
        "{",
        "    enqueue(&datas, present_dat);",
        "    enqueue(&times, present_time);",
        "    if (/* local max on PPG window */)",
        "        enqueue(&HR_List, 60000/(t0 - t1));",
        "    return Hr_Ave_Filter(HR_List.data, 7);",
        "}",
    ]
    p1 = FIG_DIR / "图4-3_心率算法核心片段.png"
    code_lines_to_png(hr_lines, p1, title="HrAlgorythm.c (excerpt)")
    p2 = FIG_DIR / "图2-1_系统总体框图示意.png"
    draw_block_diagram(p2)
    return p1, p2


def add_table_tasks(doc):
    add_fig_note(doc, "【插图4-2】可选：Keil 中 user_TasksInit.c 任务创建段落截屏（与表4-1二选一或并存）。")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "任务名"
    hdr[1].text = "作用（我用白话写）"
    hdr[2].text = "备注"
    rows = [
        ("HardwareInitTask", "上电只做一次的硬件与 LVGL 初始化", "跑完自删"),
        ("LvHandlerTask", "周期性调用 lv_timer_handler", "界面心跳"),
        ("WDOGFeedTask", "喂看门狗", "防死机"),
        ("IdleEnterTask / StopEnterTask", "配合熄屏、停止模式", "省电相关"),
        ("KeyTask", "按键消息", "事件进队列"),
        ("ScrRenewTask", "页面切换与刷新调度", "和 UI 页面绑定"),
        ("SensorDataUpdateTask", "温湿度、指南针、环境等数据更新", "按页面开门测量"),
        ("HRDataUpdateTask", "心率页打开时采 PPG 并算心率", "50 ms 节拍附近"),
        ("DataSaveTask", "设置、步数等落 EEPROM", "小数据块"),
        ("MessageSendTask", "串口/蓝牙侧指令与回传", "调试与扩展"),
        ("MPUCheckTask", "抬腕/姿态粗判", "可触发熄屏策略"),
    ]
    for name, role, note in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = role
        row[2].text = note
    p = doc.add_paragraph("表4-1  FreeRTOS 任务分工一览")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def count_chinese_chars(doc):
    s = "\n".join(p.text for p in doc.paragraphs)
    return len(re.findall(r"[\u4e00-\u9fff]", s))


def add_chapter_from_dict(doc, chapter_title, chapter_dict, level=1):
    doc.add_heading(chapter_title, level=level)
    for sec_title, paras in chapter_dict.items():
        doc.add_heading(sec_title, level=2)
        for t in paras:
            add_body(doc, t)


def main():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig_hr, fig_sys = build_figures()

    doc = Document()
    set_doc_default_font(doc)

    # ----- 摘要 -----
    doc.add_heading("摘  要", level=1)
    for t in 摘要段落:
        add_body(doc, t, first_indent=True)
    add_body(doc, 关键词行, first_indent=False)

    doc.add_heading("ABSTRACT", level=1)
    add_body(
        doc,
        "Battery life strongly affects the usability of smart watches. This project builds a wearable prototype "
        "that balances basic functions and power saving. An STM32F411 MCU runs FreeRTOS and LVGL, with common "
        "sensors and a Bluetooth module for future phone link. Sleep and selective sensor power-on are used to "
        "cut idle cost. Lab tests show smooth UI and stable data paths; endurance depends on battery and usage, "
        "and measurement methods are described.",
        first_indent=False,
    )
    add_body(
        doc,
        "Keywords: low power; smart watch; STM32; FreeRTOS; LVGL; wearable",
        first_indent=False,
    )

    # ----- 第1章 -----
    add_chapter_from_dict(doc, "第1章  绪论", 第1章)
    add_fig_note(doc, "【插图1-1】可选：市面常见智能手表/手环产品拼图或报道截图（注意引用的规范性）。")

    # ----- 第2章 -----
    doc.add_heading("第2章  总体方案与相关技术", level=1)
    for sec_title, paras in 第2章.items():
        if sec_title.startswith("2.2"):
            doc.add_heading(sec_title, level=2)
            for t in paras:
                add_body(doc, t)
            add_fig_note(doc, "【插图2-1】系统总体框图（下方为脚本生成的示意图，可换成正式框图）。")
            doc.add_picture(str(fig_sys), width=Inches(5.8))
            p = doc.add_paragraph("图2-1  系统总体结构示意")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if sec_title.startswith("2.3"):
            doc.add_heading(sec_title, level=2)
            for t in paras:
                add_body(doc, t)
            add_fig_note(doc, "【插图2-2】Keil 工程配置截图（芯片 STM32F411CEUx，宏 STM32F411xE）。")
            continue
        doc.add_heading(sec_title, level=2)
        for t in paras:
            add_body(doc, t)

    # ----- 第3章 -----
    doc.add_heading("第3章  硬件设计", level=1)
    for sec_title, paras in 第3章.items():
        doc.add_heading(sec_title, level=2)
        for t in paras:
            add_body(doc, t)
        if sec_title.startswith("3.1"):
            add_fig_note(doc, "【插图3-1】电源与主控原理图截图。")
        elif sec_title.startswith("3.2"):
            add_fig_note(doc, "【插图3-2】屏幕与触摸模块实物或接口照片。")
        elif sec_title.startswith("3.3"):
            add_fig_note(doc, "【插图3-3】传感器在 PCB 上位置标注或装配图。")

    # ----- 第4章 -----
    doc.add_heading("第4章  软件设计", level=1)
    sec43_paras = None
    for sec_title, paras in 第4章.items():
        if sec_title.startswith("4.3"):
            sec43_paras = paras
            continue
        if sec_title.startswith("4.4"):
            doc.add_heading(sec_title, level=2)
            for t in paras:
                add_body(doc, t)
            add_fig_note(doc, "图4-3  心率估算核心代码片段（脚本生成截图）")
            doc.add_picture(str(fig_hr), width=Inches(5.8))
            continue
        doc.add_heading(sec_title, level=2)
        for t in paras:
            add_body(doc, t)
        if sec_title.startswith("4.2"):
            add_fig_note(doc, "【插图4-1】主页、菜单等实机照片。")

    doc.add_heading("4.3  任务与通信", level=2)
    if sec43_paras:
        add_body(doc, sec43_paras[0])
    add_table_tasks(doc)
    if sec43_paras and len(sec43_paras) > 1:
        for t in sec43_paras[1:]:
            add_body(doc, t)

    # ----- 第5章 -----
    doc.add_heading("第5章  测试与结果", level=1)
    for sec_title, paras in 第5章.items():
        doc.add_heading(sec_title, level=2)
        for t in paras:
            add_body(doc, t)
        if sec_title.startswith("5.1"):
            t = doc.add_table(rows=1, cols=4)
            t.style = "Table Grid"
            h = t.rows[0].cells
            h[0].text = "项目"
            h[1].text = "操作步骤（简述）"
            h[2].text = "期望结果"
            h[3].text = "结果记录"
            for a, b, c, d in [
                ("时间显示", "上电观察主页", "与 RTC 一致", "□通过"),
                ("触摸", "滑动、点击各页", "页面切换正常", "□通过"),
                ("温湿度", "进入环境页", "数值合理变化", "□通过"),
                ("心率", "佩戴进入心率页", "出现合理读数", "□通过"),
                ("指南针", "进入指南针页缓慢旋转", "方向变化平滑", "□通过"),
                ("蓝牙/串口", "发测试指令", "有正确应答", "□通过"),
                ("充电指示", "连接充电器", "有充电图标或提示", "□通过"),
            ]:
                r = t.add_row().cells
                r[0].text, r[1].text, r[2].text, r[3].text = a, b, c, d
            p = doc.add_paragraph("表5-1  主要功能测试检查表（□处打勾或写说明）")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_fig_note(doc, "【插图5-1】测试过程拍照拼图（建议带时间与固件版本备注）。")

    # ----- 结论 -----
    doc.add_heading("结  论", level=1)
    for t in 结论段落:
        add_body(doc, t)

    # ----- 参考文献 -----
    doc.add_heading("参考文献", level=1)
    refs = [
        "意法半导体. STM32F411xC/E 系列参考手册[EB/OL].",
        "意法半导体. STM32Cube HAL 驱动说明 UM1850[EB/OL].",
        "Richard Barry. Mastering the FreeRTOS Real Time Kernel（内核说明与示例）[EB/OL].",
        "LVGL 官方文档：图形库配置、控件与字体裁剪指南[EB/OL]. https://docs.lvgl.io/",
        "ARM Limited. Cortex-M4 技术参考手册（指令集与异常模型查阅）[EB/OL].",
        "开源硬件社区. OV-Watch 等圆形智能手表开源资料（硬件与软件参考）[EB/OL].",
        "陈国良. 嵌入式实时操作系统在消费电子中的典型应用综述类论文（请用知网实条替换作者与刊名）[J].",
        "刘洋等. 可穿戴设备低功耗设计相关中文论文（请用知网实条替换）[J].",
        "张建等. 基于 STM32 的物联网终端设计类参考文献（请用知网实条替换）[J].",
        "王磊. 触摸屏与 SPI 显示屏驱动实践类教材或论文（请按学校格式补全）[M].",
    ]
    for i, r in enumerate(refs, 1):
        doc.add_paragraph(f"[{i}] {r}", style="List Number")

    # ----- 附录 -----
    doc.add_heading("附  录", level=1)
    for t in 附录段落:
        add_body(doc, t)
    add_fig_note(doc, "【插图附-1】整机装配完成图（建议高清）。")

    # ----- 致谢 -----
    doc.add_heading("致  谢", level=1)
    for t in 致谢段落:
        add_body(doc, t)

    doc.save(OUT_DOCX)
    zh_n = count_chinese_chars(doc)
    print("Wrote:", OUT_DOCX)
    print("Figures:", fig_hr, fig_sys)
    print("Approx. Chinese characters in docx (all paragraphs):", zh_n)


if __name__ == "__main__":
    main()
